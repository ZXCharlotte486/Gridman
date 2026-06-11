"""
函证生成工具

输入：函证对象清单（来自审计抽样输出 + 客户/供应商基础信息）
功能：批量生成函证 Word 文件 + 函证清单 Excel

支持的函证类型：
- bank：银行询证函（询余额、借款、未结算业务、保证金等）
- receivable：应收账款询证函（询客户欠款余额）
- payable：应付账款询证函（询供应商欠款余额）
- lawyer：律师函证（询诉讼/或有事项）

设计原则：古立特只生成函证的"核心内容"——对方信息 + 询证事项 + 数据。
封面、抬头、签字栏、装订格式由各所自己处理（古立特不做模板适配）。
"""
from pathlib import Path
from typing import Optional
from datetime import datetime


# 各类型函证的字段定义
CONFIRMATION_FIELDS = {
    "bank": {
        "name": "银行存款及借款询证函",
        "subject_label": "贵银行",
        "key_fields": [
            "账号",
            "币种",
            "存款余额",
            "借款余额",
            "未结算业务",
            "保证金",
            "授信额度",
        ],
    },
    "receivable": {
        "name": "应收账款询证函",
        "subject_label": "贵公司",
        "key_fields": [
            "应收账款余额",
            "币种",
        ],
    },
    "payable": {
        "name": "应付账款询证函",
        "subject_label": "贵公司",
        "key_fields": [
            "应付账款余额",
            "币种",
        ],
    },
    "lawyer": {
        "name": "律师函证",
        "subject_label": "贵律师事务所",
        "key_fields": [
            "案件名称",
            "诉讼标的",
            "案件状态",
        ],
    },
}


def generate_confirmation_letters(
    sample_path: str,
    confirmation_type: str,
    entity_name: str,
    cutoff_date: str,
    return_address: str = "",
    return_email: str = "",
    counterparty_column: str = "对方名称",
    amount_column: str = "金额",
    address_column: str = "地址",
    contact_column: str = "联系人",
    extra_columns: Optional[list] = None,
    sheet_name: Optional[str] = None,
    output_dir: Optional[str] = None,
) -> dict:
    """
    批量生成函证 Word 文件 + 函证清单 Excel。

    Args:
        sample_path: 函证样本表（Excel/CSV，通常是 audit_sampling 的输出）
        confirmation_type: 函证类型 bank / receivable / payable / lawyer
        entity_name: 被审计单位名称
        cutoff_date: 函证截止日 YYYY-MM-DD
        return_address: 回函地址（事务所地址）
        return_email: 回函邮箱（如有）
        counterparty_column: 对方名称列名
        amount_column: 金额列名
        address_column: 地址列名（可选）
        contact_column: 联系人列名（可选）
        extra_columns: 额外要写入函证的列（如银行函证的"账号"列）
        sheet_name: Excel 工作表名
        output_dir: 输出目录（默认 ./函证_YYYYMMDD/）

    Returns:
        dict: {
            "status": "success"|"error",
            "summary": {
                "函证总数": N,
                "函证类型": str,
                "Word 文件": str,
                "Excel 清单": str,
            },
            "letters": [{编号, 对方, 金额, 状态, ...}],
            "output_dir": str
        }
    """
    try:
        import pandas as pd

        # python-docx 为可选依赖（审计 extra），缺失时友好降级
        try:
            import docx  # noqa: F401
        except ImportError:
            return {
                "status": "error",
                "message": "python-docx 未安装（生成 Word 函证为审计可选功能）。"
                           "启用方式：将 uvx 配置的 --from 改为 \"gridman-mcp[audit]\"，"
                           "或在当前环境执行 pip install python-docx",
            }

        # 校验类型
        if confirmation_type not in CONFIRMATION_FIELDS:
            return {
                "status": "error",
                "message": f"不支持的函证类型：{confirmation_type}，"
                           f"可选：{list(CONFIRMATION_FIELDS.keys())}",
            }

        # 读取
        fp = Path(sample_path)
        if not fp.exists():
            return {"status": "error", "message": f"文件不存在：{sample_path}"}

        if fp.suffix.lower() in (".xlsx", ".xls"):
            df = pd.read_excel(fp, sheet_name=sheet_name or 0)
        elif fp.suffix.lower() == ".csv":
            df = pd.read_csv(fp)
        else:
            return {"status": "error", "message": f"不支持的格式：{fp.suffix}"}

        # 验证必需列
        if counterparty_column not in df.columns:
            return {
                "status": "error",
                "message": f"缺少列 '{counterparty_column}'，可用列：{list(df.columns)}",
            }

        # 校验截止日
        try:
            cutoff = datetime.strptime(cutoff_date, "%Y-%m-%d").date()
        except ValueError:
            return {"status": "error", "message": f"cutoff_date 格式错误：{cutoff_date}"}

        # 准备输出目录
        if output_dir is None:
            output_dir = f"函证_{cutoff.strftime('%Y%m%d')}"
        out_dir = Path(output_dir)
        out_dir.mkdir(parents=True, exist_ok=True)

        type_info = CONFIRMATION_FIELDS[confirmation_type]

        # ── 1. 生成 Word 文件 ──
        word_path = out_dir / f"{type_info['name']}_批量.docx"
        letters_meta = _generate_word_letters(
            df=df,
            confirmation_type=confirmation_type,
            entity_name=entity_name,
            cutoff_date=cutoff,
            return_address=return_address,
            return_email=return_email,
            counterparty_column=counterparty_column,
            amount_column=amount_column,
            address_column=address_column,
            contact_column=contact_column,
            extra_columns=extra_columns or [],
            output_path=word_path,
        )

        # ── 2. 生成 Excel 清单 ──
        excel_path = out_dir / f"{type_info['name']}_清单.xlsx"
        _generate_excel_list(
            letters_meta=letters_meta,
            confirmation_type=confirmation_type,
            entity_name=entity_name,
            cutoff_date=cutoff,
            output_path=excel_path,
        )

        return {
            "status": "success",
            "summary": {
                "函证类型": type_info["name"],
                "函证总数": len(letters_meta),
                "被审计单位": entity_name,
                "截止日": cutoff_date,
                "Word 文件": str(word_path),
                "Excel 清单": str(excel_path),
            },
            "letters": letters_meta[:50],
            "output_dir": str(out_dir),
            "message": (
                f"已生成 {len(letters_meta)} 份{type_info['name']}：\n"
                f"  Word（每对象一页）: {word_path.name}\n"
                f"  Excel 函证清单:    {excel_path.name}"
            ),
        }

    except ImportError as e:
        return {"status": "error", "message": f"依赖未安装：{e}（需要 python-docx）"}
    except Exception as e:
        import traceback
        return {"status": "error", "message": str(e), "trace": traceback.format_exc()}


# ── Word 文件生成 ──

def _generate_word_letters(
    df, confirmation_type, entity_name, cutoff_date,
    return_address, return_email,
    counterparty_column, amount_column, address_column, contact_column,
    extra_columns, output_path,
):
    """生成 Word 函证（每对象一页）"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn

    type_info = CONFIRMATION_FIELDS[confirmation_type]
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    letters_meta = []

    for idx, row in df.iterrows():
        counterparty = str(row[counterparty_column]).strip()
        if not counterparty or counterparty == "nan":
            continue

        letter_no = f"{confirmation_type.upper()}-{idx + 1:04d}"
        amount = 0.0
        if amount_column in df.columns:
            try:
                v = row[amount_column]
                amount = float(v) if v is not None and str(v) != "nan" else 0.0
            except (ValueError, TypeError):
                amount = 0.0

        address = str(row[address_column]).strip() if address_column in df.columns and str(row.get(address_column, "")) != "nan" else ""
        contact = str(row[contact_column]).strip() if contact_column in df.columns and str(row.get(contact_column, "")) != "nan" else ""

        # 标题（居中加粗）
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(type_info["name"])
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(0x2D, 0x2B, 0x55)

        # 函证编号
        no_para = doc.add_paragraph()
        no_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        no_run = no_para.add_run(f"函证编号：{letter_no}")
        no_run.font.size = Pt(10)

        # 收件方
        recipient_para = doc.add_paragraph()
        recipient_run = recipient_para.add_run(f"{counterparty}：")
        recipient_run.bold = True
        recipient_run.font.size = Pt(12)

        # 主体段落（按类型分支）
        body_paragraphs = _build_body(
            confirmation_type=confirmation_type,
            entity_name=entity_name,
            counterparty=counterparty,
            amount=amount,
            cutoff_date=cutoff_date,
            row=row,
            extra_columns=extra_columns,
        )
        for text in body_paragraphs:
            p = doc.add_paragraph(text)
            p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进 2 字符
            p.paragraph_format.line_spacing = 1.5

        # 数据表格（银行函证按14项标准格式）
        if confirmation_type == "bank":
            doc.add_paragraph()
            # 第1项：银行存款
            p1 = doc.add_paragraph()
            p1.add_run("1. 银行存款").bold = True
            tbl = doc.add_table(rows=2, cols=6)
            tbl.style = "Light Grid Accent 1"
            headers = ["账户名称", "银行账号", "币种", "账户类型", "账户余额", "是否存在冻结/担保/受限"]
            for i, h in enumerate(headers):
                tbl.rows[0].cells[i].text = h
            # 填充数据
            tbl.rows[1].cells[0].text = entity_name
            for col in extra_columns:
                if col in df.columns:
                    val = str(row.get(col, ""))
                    if val == "nan":
                        val = ""
                    # 尝试匹配到对应列
                    for i, h in enumerate(headers):
                        if col in h or h in col:
                            tbl.rows[1].cells[i].text = val
                            break
            # 金额填入余额列
            tbl.rows[1].cells[4].text = f"{amount:,.2f}"

            doc.add_paragraph()
            # 第2项：银行借款
            p2 = doc.add_paragraph()
            p2.add_run("2. 银行借款").bold = True
            tbl2 = doc.add_table(rows=2, cols=5)
            tbl2.style = "Light Grid Accent 1"
            headers2 = ["借款账号", "币种", "余额", "到期日期", "利率"]
            for i, h in enumerate(headers2):
                tbl2.rows[0].cells[i].text = h
            tbl2.rows[1].cells[0].text = "无"

            doc.add_paragraph()
            # 第3-14项简化（标注"无"或划线）
            other_items = [
                "3. 注销的银行存款账户",
                "4. 本公司作为委托人的委托贷款",
                "5. 本公司作为借款人的委托贷款",
                "6. 担保",
                "7. 银行承兑汇票",
                "8. 已贴现未到期商业汇票",
                "9. 托收/提示付款的商业汇票",
                "10. 不可撤销信用证",
                "11. 外汇买卖合约",
                "12. 托管的证券或产权文件",
                "13. 未到期银行理财产品",
                "14. 其他",
            ]
            for item in other_items:
                pi = doc.add_paragraph()
                pi.add_run(item).bold = True
                pi.paragraph_format.space_after = Pt(2)
                pv = doc.add_paragraph("    无")
                pv.paragraph_format.space_after = Pt(6)
        else:
            # 非银行函证：简单余额展示
            balance_para = doc.add_paragraph()
            balance_para.paragraph_format.first_line_indent = Cm(0.74)
            run = balance_para.add_run(f"截至 {cutoff_date.strftime('%Y年%m月%d日')}，金额为：")
            run.font.size = Pt(11)
            amt_run = balance_para.add_run(f" ¥ {amount:,.2f}")
            amt_run.bold = True
            amt_run.font.color.rgb = RGBColor(0x2D, 0x2B, 0x55)

        # 回函要求（银行函证按标准格式）
        doc.add_paragraph()
        if confirmation_type == "bank":
            return_para = doc.add_paragraph()
            return_para.paragraph_format.first_line_indent = Cm(0.74)
            return_para.add_run(
                "本公司谨授权贵行将回函直接寄至会计师事务所，地址及联系方式如下："
            )
        else:
            return_para = doc.add_paragraph()
            return_para.paragraph_format.first_line_indent = Cm(0.74)
            return_para.add_run(
                "请贵方核对上述内容是否相符，如相符请于回函栏内签章；"
                "如不符请将贵方实际金额填列于回函栏内并签章证明，直接寄回我所审计人员。"
            )

        if return_address:
            addr_para = doc.add_paragraph()
            addr_para.paragraph_format.first_line_indent = Cm(0.74)
            addr_para.add_run(f"回函地址：{return_address}")
        if return_email:
            email_para = doc.add_paragraph()
            email_para.paragraph_format.first_line_indent = Cm(0.74)
            email_para.add_run(f"回函邮箱：{return_email}")

        # 回函栏
        doc.add_paragraph()
        reply_para = doc.add_paragraph()
        if confirmation_type == "bank":
            reply_run = reply_para.add_run("【以下由被询证银行填列】")
        else:
            reply_run = reply_para.add_run("【回函栏】")
        reply_run.bold = True
        reply_para.paragraph_format.line_spacing = 2

        if confirmation_type == "bank":
            # 银行函证标准回函格式
            reply_table = doc.add_table(rows=4, cols=2)
            reply_table.style = "Light Grid Accent 1"
            reply_table.rows[0].cells[0].text = "经本行核对，所函证项目与本行记载信息相符。"
            reply_table.rows[0].cells[1].text = ""
            reply_table.rows[1].cells[0].text = "经本行核对，所函证项目存在以下不符之处："
            reply_table.rows[1].cells[1].text = ""
            reply_table.rows[2].cells[0].text = "经办人：      职务/岗位：      电话："
            reply_table.rows[2].cells[1].text = "复核人：      职务/岗位：      电话："
            reply_table.rows[3].cells[0].text = "（银行盖章）"
            reply_table.rows[3].cells[1].text = "年    月    日"
        else:
            reply_table = doc.add_table(rows=3, cols=2)
            reply_table.style = "Light Grid Accent 1"
            reply_table.rows[0].cells[0].text = "□ 上述内容相符"
            reply_table.rows[0].cells[1].text = ""
            reply_table.rows[1].cells[0].text = "□ 不相符，实际金额为"
            reply_table.rows[1].cells[1].text = ""
            reply_table.rows[2].cells[0].text = "回函单位（盖章）"
            reply_table.rows[2].cells[1].text = "回函日期：    年    月    日"

        # 落款
        doc.add_paragraph()
        signer_para = doc.add_paragraph()
        signer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signer_para.add_run(f"{entity_name}\n")
        signer_para.add_run(f"{cutoff_date.strftime('%Y年%m月%d日')}")

        # 分页
        doc.add_page_break()

        letters_meta.append({
            "函证编号": letter_no,
            "对方名称": counterparty,
            "金额": round(amount, 2),
            "地址": address,
            "联系人": contact,
            "状态": "待发函",
            "回函日期": "",
            "回函结果": "",
        })

    doc.save(str(output_path))
    return letters_meta


def _build_body(
    confirmation_type, entity_name, counterparty, amount, cutoff_date, row, extra_columns
):
    """根据函证类型生成主体段落"""
    cutoff_str = cutoff_date.strftime("%Y年%m月%d日")

    if confirmation_type == "bank":
        return [
            f"本公司聘请的会计师事务所正在对本公司（{entity_name}）"
            f"{cutoff_str}的财务报表进行审计，按照中国注册会计师审计准则的要求，"
            f"应当询证本公司与贵行相关的信息。",
            "下列第1—14项及附表（如适用）信息出自本公司的记录：",
            '（1）如与贵行记录相符，请在本函\u201c结论\u201d部分签章；',
            '（2）如有不符，请在本函\u201c结论\u201d部分列明不符项目及具体内容，并签章。',
            "本公司谨授权贵行将回函直接寄至会计师事务所。",
        ]

    elif confirmation_type == "receivable":
        return [
            f"为审计需要，现将本公司（{entity_name}）截至 {cutoff_str} "
            f"对贵公司的应收账款余额情况向贵方函证。",
            f"敬请贵公司核对相关账目，如有不符请将贵方实际欠款金额填明并寄回。",
        ]

    elif confirmation_type == "payable":
        return [
            f"为审计需要，现将本公司（{entity_name}）截至 {cutoff_str} "
            f"对贵公司的应付账款余额情况向贵方函证。",
            f"敬请贵公司核对相关账目，如有不符请将贵方实际应收款金额填明并寄回。",
        ]

    elif confirmation_type == "lawyer":
        return [
            f"为审计需要，现就本公司（{entity_name}）委托贵律师事务所"
            f"代理的诉讼、仲裁及其他法律事项情况进行函证。",
            f"截至 {cutoff_str}，请贵所提供以下信息：",
            f"（1）正在代理或已经办结的所有案件清单；",
            f"（2）每个案件的当事人、案由、诉讼标的、当前进展；",
            f"（3）律师对案件最终结果的初步判断（如能提供）。",
        ]

    return [f"为审计需要，请贵方核对相关账目情况。"]


# ── Excel 清单生成（古立特统一配色） ──

def _generate_excel_list(letters_meta, confirmation_type, entity_name, cutoff_date, output_path):
    """生成函证清单 Excel"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    from openpyxl.utils import get_column_letter

    HEADER_BG = "2D2B55"
    TOTAL_BG = "3D3A6B"
    SUBTOTAL_BG = "EEEAF5"
    ALT_BG = "F7F5FA"
    BORDER_COLOR = "D4D0E0"
    HEADER_TEXT = "E8E6F0"
    TEXT_COLOR = "2D2D3F"

    title_font = Font(name="微软雅黑", bold=True, size=14, color=HEADER_BG)
    header_font = Font(name="微软雅黑", bold=True, size=10.5, color=HEADER_TEXT)
    normal_font = Font(name="微软雅黑", size=10.5, color=TEXT_COLOR)
    total_font = Font(name="微软雅黑", bold=True, size=11, color=HEADER_TEXT)

    header_fill = PatternFill(start_color=HEADER_BG, end_color=HEADER_BG, fill_type="solid")
    alt_fill = PatternFill(start_color=ALT_BG, end_color=ALT_BG, fill_type="solid")
    subtotal_fill = PatternFill(start_color=SUBTOTAL_BG, end_color=SUBTOTAL_BG, fill_type="solid")
    total_fill = PatternFill(start_color=TOTAL_BG, end_color=TOTAL_BG, fill_type="solid")

    border = Border(
        left=Side(style="thin", color=BORDER_COLOR),
        right=Side(style="thin", color=BORDER_COLOR),
        top=Side(style="thin", color=BORDER_COLOR),
        bottom=Side(style="thin", color=BORDER_COLOR),
    )
    center = Alignment(horizontal="center", vertical="center")
    left_align = Alignment(horizontal="left", vertical="center", indent=1)
    right_align = Alignment(horizontal="right", vertical="center", indent=1)

    wb = Workbook()
    ws = wb.active
    ws.title = "函证清单"

    type_info = CONFIRMATION_FIELDS[confirmation_type]

    # 标题
    cols = ["函证编号", "对方名称", "金额", "地址", "联系人", "状态", "回函日期", "回函结果"]
    ncols = len(cols)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncols)
    ws.cell(row=1, column=1, value=f"{type_info['name']} — {entity_name}").font = title_font
    ws.cell(row=1, column=1).alignment = center
    ws.row_dimensions[1].height = 28

    # 副标题
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=ncols)
    sub = ws.cell(row=2, column=1, value=f"函证截止日：{cutoff_date.strftime('%Y-%m-%d')}    "
                                          f"共 {len(letters_meta)} 份")
    sub.font = normal_font
    sub.alignment = center

    # 表头
    for c_idx, col_name in enumerate(cols, start=1):
        cell = ws.cell(row=3, column=c_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border

    # 数据
    for r_idx, item in enumerate(letters_meta, start=4):
        for c_idx, col_name in enumerate(cols, start=1):
            val = item.get(col_name, "")
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.font = normal_font
            cell.border = border

            if col_name == "金额":
                cell.alignment = right_align
                if isinstance(val, (int, float)):
                    cell.number_format = "#,##0.00"
            elif col_name in ("函证编号", "状态", "回函日期"):
                cell.alignment = center
            else:
                cell.alignment = left_align

            if (r_idx - 4) % 2 == 1:
                cell.fill = alt_fill

    # 合计行
    total_row = 4 + len(letters_meta)
    ws.merge_cells(start_row=total_row, start_column=1, end_row=total_row, end_column=2)
    ws.cell(row=total_row, column=1, value="合计").font = total_font
    ws.cell(row=total_row, column=1).fill = total_fill
    ws.cell(row=total_row, column=1).alignment = center
    ws.cell(row=total_row, column=1).border = border
    ws.cell(row=total_row, column=2).fill = total_fill
    ws.cell(row=total_row, column=2).border = border

    total_amount = sum(item.get("金额", 0) for item in letters_meta if isinstance(item.get("金额"), (int, float)))
    cell = ws.cell(row=total_row, column=3, value=round(total_amount, 2))
    cell.font = total_font
    cell.fill = total_fill
    cell.alignment = right_align
    cell.border = border
    cell.number_format = "#,##0.00"

    for c in range(4, ncols + 1):
        ws.cell(row=total_row, column=c).fill = total_fill
        ws.cell(row=total_row, column=c).border = border

    # 列宽
    widths = [14, 28, 16, 28, 12, 10, 13, 16]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    ws.row_dimensions[3].height = 22
    ws.freeze_panes = "C4"

    wb.save(str(output_path))
