"""
审计底稿初始化工具

输入：客户名称 + 审计年度
输出：完整的底稿目录结构（按手册标准 7 大文件夹）+ 空白模板文件

设计原则：
- 完全照抄《审计底稿全流程操作手册》第八部分的目录结构
- 模板用古立特灰紫配色（用户可自行修改）
- 各所封面/抬头/装订格式由用户自行处理
- 文档类（评价表/总结/报告）→ Word，表格类（计划/明细/调整）→ Excel
"""
from pathlib import Path
from datetime import datetime
from typing import Optional

from gridman_mcp import __version__ as _GV

# 古立特统一配色
HEADER_BG = "2D2B55"
TOTAL_BG = "3D3A6B"
SUBTOTAL_BG = "EEEAF5"
ALT_BG = "F7F5FA"
BORDER_COLOR = "D4D0E0"
HEADER_TEXT = "E8E6F0"
TEXT_COLOR = "2D2D3F"


def init_workpaper(
    client_name: str,
    audit_year: int,
    output_dir: Optional[str] = None,
    is_group_audit: bool = False,
    perform_control_test: bool = False,
) -> dict:
    """
    初始化审计底稿目录结构 + 生成空白模板。

    Args:
        client_name: 客户名称（如"某某科技股份有限公司"）
        audit_year: 审计年度（如 2025，对应 2025 年度审计）
        output_dir: 输出根目录（默认 ./{客户名}_{年度}审计底稿/）
        is_group_audit: 是否集团审计（影响 1-计划 下的模板内容）
        perform_control_test: 是否执行控制测试（影响 2-控制测试 文件夹）

    Returns:
        dict: {
            "status": "success"|"error",
            "summary": {
                "客户": str,
                "年度": int,
                "底稿根目录": str,
                "创建文件夹数": N,
                "生成模板数": M,
            },
            "structure": {目录树},
            "templates": [模板文件清单],
            "next_steps": [下一步建议]
        }
    """
    try:
        # python-docx 为可选依赖（审计 extra），缺失时友好降级
        try:
            import docx  # noqa: F401
        except ImportError:
            return {
                "status": "error",
                "message": "python-docx 未安装（生成 Word 底稿模板为审计可选功能）。"
                           "启用方式：将 uvx 配置的 --from 改为 \"gridman-mcp[audit]\"，"
                           "或在当前环境执行 pip install python-docx",
            }

        # 准备根目录
        folder_name = f"{client_name}_{audit_year}年度审计底稿"
        if output_dir is None:
            root = Path(folder_name)
        else:
            root = Path(output_dir) / folder_name
        root.mkdir(parents=True, exist_ok=True)

        # 创建七大文件夹
        folders = [
            ("1-计划与风险", True),
            ("2-控制测试", perform_control_test),
            ("3-实质性程序", True),
            ("4-准则合规", True),
            ("5-职业判断", True),
            ("6-完成阶段", True),
            ("7-备查类", True),
        ]
        created_folders = []
        for name, should_create in folders:
            if should_create:
                folder = root / name
                folder.mkdir(exist_ok=True)
                created_folders.append(name)

        # 3-实质性程序 下的子文件夹
        subprogram_folders = [
            "货币资金",
            "函证",
            "往来款项",
            "在建工程",
            "损益类",
        ]
        for sub in subprogram_folders:
            (root / "3-实质性程序" / sub).mkdir(exist_ok=True)
            created_folders.append(f"3-实质性程序/{sub}")

        # 7-备查类 下的子文件夹
        backup_folders = ["重要合同", "其他支持性文件"]
        for sub in backup_folders:
            (root / "7-备查类" / sub).mkdir(exist_ok=True)
            created_folders.append(f"7-备查类/{sub}")

        # 生成模板
        templates_created = []

        # ── 1-计划与风险 ──
        _create_business_acceptance_doc(
            root / "1-计划与风险" / "业务承接评价表.docx",
            client_name, audit_year, is_group_audit
        )
        templates_created.append("1-计划与风险/业务承接评价表.docx")

        _create_audit_strategy_doc(
            root / "1-计划与风险" / "总体审计策略.docx",
            client_name, audit_year, is_group_audit
        )
        templates_created.append("1-计划与风险/总体审计策略.docx")

        _create_specific_plan_xlsx(
            root / "1-计划与风险" / "具体审计计划汇总表.xlsx",
            client_name, audit_year
        )
        templates_created.append("1-计划与风险/具体审计计划汇总表.xlsx")

        _create_risk_assessment_xlsx(
            root / "1-计划与风险" / "风险评估结果汇总表.xlsx",
            client_name, audit_year
        )
        templates_created.append("1-计划与风险/风险评估结果汇总表.xlsx")

        if is_group_audit:
            _create_group_materiality_xlsx(
                root / "1-计划与风险" / "集团审计重要性分配表.xlsx",
                client_name, audit_year
            )
            templates_created.append("1-计划与风险/集团审计重要性分配表.xlsx")

        # ── 4-准则合规 ──
        _create_automation_log_xlsx(
            root / "4-准则合规" / "自动化工具使用记录.xlsx",
            client_name, audit_year
        )
        templates_created.append("4-准则合规/自动化工具使用记录.xlsx")

        # ── 5-职业判断 ──
        _create_judgment_xlsx(
            root / "5-职业判断" / "重大职业判断记录单.xlsx",
            client_name, audit_year
        )
        templates_created.append("5-职业判断/重大职业判断记录单.xlsx")

        # ── 6-完成阶段 ──
        _create_uncorrected_misstatement_xlsx(
            root / "6-完成阶段" / "未更正错报汇总表.xlsx",
            client_name, audit_year
        )
        templates_created.append("6-完成阶段/未更正错报汇总表.xlsx")

        _create_review_log_xlsx(
            root / "6-完成阶段" / "复核记录表.xlsx",
            client_name, audit_year
        )
        templates_created.append("6-完成阶段/复核记录表.xlsx")

        _create_post_event_xlsx(
            root / "6-完成阶段" / "期后事项复核表.xlsx",
            client_name, audit_year
        )
        templates_created.append("6-完成阶段/期后事项复核表.xlsx")

        _create_audit_summary_doc(
            root / "6-完成阶段" / "审计总结.docx",
            client_name, audit_year
        )
        templates_created.append("6-完成阶段/审计总结.docx")

        _create_audit_report_doc(
            root / "6-完成阶段" / "审计报告（草稿）.docx",
            client_name, audit_year
        )
        templates_created.append("6-完成阶段/审计报告（草稿）.docx")

        # ── 根目录的索引 ──
        _create_workpaper_index(
            root / "底稿索引.xlsx",
            client_name, audit_year, created_folders, templates_created
        )
        templates_created.append("底稿索引.xlsx")

        # ── README 配色说明 ──
        _create_readme(
            root / "_README_配色与使用说明.txt",
            client_name, audit_year
        )
        templates_created.append("_README_配色与使用说明.txt")

        next_steps = [
            "1. 用 materiality_calculator 计算重要性水平，填入 总体审计策略.docx",
            "2. 用 data_file_overview 查看科目余额表和序时账，了解数据结构",
            "3. 按 W13（底稿全流程）工作流逐步执行各项审计程序",
            "4. 各程序产出的 Excel 直接放到对应的 3-实质性程序/[科目]/ 子目录",
            "5. 完成后用 audit_adjustment 生成三栏 TB，放到 6-完成阶段/",
            "6. 古立特生成的所有 Excel 配色可在事务所内部统一修改",
        ]

        return {
            "status": "success",
            "summary": {
                "客户": client_name,
                "审计年度": audit_year,
                "集团审计": is_group_audit,
                "底稿根目录": str(root.absolute()),
                "创建文件夹数": len(created_folders),
                "生成模板数": len(templates_created),
            },
            "folders": created_folders,
            "templates": templates_created,
            "next_steps": next_steps,
            "message": (
                f"已为【{client_name}】{audit_year}年度审计创建完整底稿框架：\n"
                f"  - {len(created_folders)} 个文件夹\n"
                f"  - {len(templates_created)} 份空白模板\n"
                f"  - 根目录：{root}"
            ),
        }

    except ImportError as e:
        return {"status": "error", "message": f"依赖未安装：{e}"}
    except Exception as e:
        import traceback
        return {"status": "error", "message": str(e), "trace": traceback.format_exc()}


# ─────────────────────────────────────────────────────────────
# Word 模板生成函数
# ─────────────────────────────────────────────────────────────

def _setup_doc_style(doc):
    """设置 Word 文档默认字体（微软雅黑）"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)


def _add_title(doc, text, size=18):
    """添加居中加粗标题"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x2D, 0x2B, 0x55)
    return p


def _add_section_heading(doc, text):
    """添加二级标题"""
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x3D, 0x3A, 0x6B)
    return p


def _add_filled_paragraph(doc, label, placeholder=""):
    """添加 '标签：[填写]' 格式的段落"""
    from docx.shared import Cm
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f"{label}：")
    run.bold = True
    p.add_run(f"{placeholder if placeholder else '[请填写]'}")


def _create_business_acceptance_doc(path, client_name, audit_year, is_group_audit):
    """业务承接评价表"""
    from docx import Document
    doc = Document()
    _setup_doc_style(doc)

    _add_title(doc, "业务承接（保持）评价表")
    doc.add_paragraph()

    # 基本信息
    doc.add_paragraph(f"客户名称：{client_name}").runs[0].bold = True
    doc.add_paragraph(f"审计年度：{audit_year}年度")
    doc.add_paragraph(f"评价日期：[请填写]")
    doc.add_paragraph()

    sections = [
        ("一、客户基本情况",
         ["客户行业及主要业务", "经营状况（盈利/亏损/规模）",
          "组织结构及股权关系", "主要客户及供应商"]),
        ("二、管理层诚信评估",
         ["管理层背景及历史", "是否有舞弊或违规历史",
          "前任审计师评价（如有）", "与监管机构的关系"]),
        ("三、独立性确认",
         ["项目组成员独立性声明", "事务所与客户的经济利益关系",
          "近亲属或重要关系人在客户处任职情况", "其他可能影响独立性的事项"]),
        ("四、专业胜任能力评估",
         ["项目组人员配备", "行业专业知识储备",
          "对客户业务复杂度的评估", "是否需要专家协助（IT/估值/法律）"]),
        ("五、资源充足性",
         ["人员档期", "时间安排", "技术工具（含古立特等自动化工具）"]),
    ]

    if is_group_audit:
        sections.append((
            "六、集团审计特别考虑",
            ["集团组织结构图", "组成部分数量及分布",
             "重要组成部分识别", "组成部分注册会计师工作能否依赖",
             "组成部分会计政策一致性"]
        ))

    for title, items in sections:
        _add_section_heading(doc, title)
        for item in items:
            _add_filled_paragraph(doc, item)
        doc.add_paragraph()

    # 结论
    _add_section_heading(doc, "评价结论")
    p = doc.add_paragraph()
    p.add_run("□ 同意承接（保持）业务\n").bold = True
    p.add_run("□ 不同意承接（保持）业务，理由：[请填写]\n").bold = True

    doc.add_paragraph()
    sig = doc.add_paragraph("项目合伙人签字：________________   日期：____年__月__日")

    doc.save(str(path))


def _create_audit_strategy_doc(path, client_name, audit_year, is_group_audit):
    """总体审计策略"""
    from docx import Document
    doc = Document()
    _setup_doc_style(doc)

    _add_title(doc, "总体审计策略")
    doc.add_paragraph()

    doc.add_paragraph(f"客户名称：{client_name}").runs[0].bold = True
    doc.add_paragraph(f"审计年度：{audit_year}年度")
    doc.add_paragraph()

    _add_section_heading(doc, "一、审计范围")
    _add_filled_paragraph(doc, "财务报表组成部分", "资产负债表 / 利润表 / 现金流量表 / 所有者权益变动表 / 附注")
    _add_filled_paragraph(doc, "审计准据", "企业会计准则 + CSA 审计准则")
    _add_filled_paragraph(doc, "本期重大变化", "[请填写]")
    doc.add_paragraph()

    _add_section_heading(doc, "二、重要性水平")
    _add_filled_paragraph(doc, "财务报表整体重要性", "[使用 materiality_calculator 工具计算]")
    _add_filled_paragraph(doc, "计算依据", "税前利润 × 5%（建议值，可调整）")
    _add_filled_paragraph(doc, "实际执行重要性", "重要性 × 60%（建议值，可调整）")
    _add_filled_paragraph(doc, "明显微小错报", "重要性 × 5%（建议值，可调整）")

    if is_group_audit:
        _add_filled_paragraph(doc, "集团重要性", "[请填写]")
        _add_filled_paragraph(doc, "组件重要性（双限值）", "≤ 0.8 × 集团重要性")
    doc.add_paragraph()

    _add_section_heading(doc, "三、审计方向")
    _add_filled_paragraph(doc, "高风险领域", "[请基于风险评估填写]")
    _add_filled_paragraph(doc, "拟应对措施", "[请填写]")
    doc.add_paragraph()

    _add_section_heading(doc, "四、时间安排")
    for stage in ["预审", "年审", "复核", "报告日"]:
        _add_filled_paragraph(doc, stage, "[起止日期]")
    doc.add_paragraph()

    _add_section_heading(doc, "五、资源安排")
    _add_filled_paragraph(doc, "项目合伙人", "")
    _add_filled_paragraph(doc, "项目经理", "")
    _add_filled_paragraph(doc, "审计员", "")
    _add_filled_paragraph(doc, "专家协助", "（如适用）")
    _add_filled_paragraph(doc, "自动化工具", f"古立特 Gridman v{_GV}（详见 4-准则合规/自动化工具使用记录.xlsx）")

    doc.save(str(path))


def _create_audit_summary_doc(path, client_name, audit_year):
    """审计总结"""
    from docx import Document
    doc = Document()
    _setup_doc_style(doc)

    _add_title(doc, "审计总结")
    doc.add_paragraph()
    doc.add_paragraph(f"客户名称：{client_name}").runs[0].bold = True
    doc.add_paragraph(f"审计年度：{audit_year}年度")
    doc.add_paragraph()

    sections = [
        "一、审计计划执行情况",
        "二、重大事项及处理",
        "三、风险评估结果更新",
        "四、调整分录及未更正错报的影响",
        "五、独立性与职业道德声明",
        "六、关键审计事项",
        "七、与治理层沟通情况",
        "八、审计意见类型及理由",
    ]
    for sec in sections:
        _add_section_heading(doc, sec)
        doc.add_paragraph("[请填写]")
        doc.add_paragraph()

    doc.save(str(path))


def _create_audit_report_doc(path, client_name, audit_year):
    """审计报告草稿"""
    from docx import Document
    doc = Document()
    _setup_doc_style(doc)

    _add_title(doc, "审 计 报 告（草稿）")
    doc.add_paragraph()
    doc.add_paragraph(f"[报告字号]")
    doc.add_paragraph()

    doc.add_paragraph(f"{client_name}全体股东：")
    doc.add_paragraph()

    sections = [
        ("一、审计意见", "我们审计了{client}财务报表，包括{date}的合并及母公司资产负债表，"
                          "{year}年度的合并及母公司利润表、合并及母公司现金流量表、"
                          "合并及母公司所有者权益变动表，以及相关财务报表附注。\n\n"
                          "我们认为，[请填写：财务报表已经/在所有重大方面/严重失实地]按照"
                          "企业会计准则的规定编制，公允反映了{client}{date}的财务状况以及"
                          "{year}年度的经营成果和现金流量。"),
        ("二、形成审计意见的基础",
         "我们按照中国注册会计师审计准则的规定执行了审计工作。"
         "审计报告的'注册会计师对财务报表审计的责任'部分进一步阐述了我们在这些准则下的责任。"
         "按照中国注册会计师职业道德守则，我们独立于{client}，并履行了职业道德方面的其他责任。"
         "我们相信，我们获取的审计证据是充分、适当的，为发表审计意见提供了基础。"),
        ("三、关键审计事项", "[请列示本年度关键审计事项及应对措施]"),
        ("四、其他信息", "[如适用]"),
        ("五、管理层对财务报表的责任", "[标准段落]"),
        ("六、注册会计师对财务报表审计的责任", "[标准段落]"),
    ]
    for title, content in sections:
        _add_section_heading(doc, title)
        text = content.replace("{client}", client_name).replace(
            "{year}", str(audit_year)).replace(
            "{date}", f"{audit_year}年12月31日")
        doc.add_paragraph(text)
        doc.add_paragraph()

    doc.add_paragraph()
    doc.add_paragraph("XX 会计师事务所（特殊普通合伙）          中国注册会计师：________________")
    doc.add_paragraph()
    doc.add_paragraph("                                          中国注册会计师：________________")
    doc.add_paragraph()
    doc.add_paragraph("中国 [城市]                              [审计报告日]")

    doc.save(str(path))


# ─────────────────────────────────────────────────────────────
# Excel 模板生成函数（古立特统一配色）
# ─────────────────────────────────────────────────────────────

def _get_styles():
    """统一样式"""
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    return {
        "title_font": Font(name="微软雅黑", bold=True, size=14, color=HEADER_BG),
        "header_font": Font(name="微软雅黑", bold=True, size=10.5, color=HEADER_TEXT),
        "normal_font": Font(name="微软雅黑", size=10.5, color=TEXT_COLOR),
        "total_font": Font(name="微软雅黑", bold=True, size=11, color=HEADER_TEXT),
        "header_fill": PatternFill(start_color=HEADER_BG, end_color=HEADER_BG, fill_type="solid"),
        "alt_fill": PatternFill(start_color=ALT_BG, end_color=ALT_BG, fill_type="solid"),
        "subtotal_fill": PatternFill(start_color=SUBTOTAL_BG, end_color=SUBTOTAL_BG, fill_type="solid"),
        "total_fill": PatternFill(start_color=TOTAL_BG, end_color=TOTAL_BG, fill_type="solid"),
        "border": Border(
            left=Side(style="thin", color=BORDER_COLOR),
            right=Side(style="thin", color=BORDER_COLOR),
            top=Side(style="thin", color=BORDER_COLOR),
            bottom=Side(style="thin", color=BORDER_COLOR),
        ),
        "center": Alignment(horizontal="center", vertical="center"),
        "left_align": Alignment(horizontal="left", vertical="center", indent=1),
    }


def _write_xlsx_template(path, sheet_name, title, headers, sample_rows, col_widths):
    """通用 xlsx 模板生成"""
    from openpyxl import Workbook
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    s = _get_styles()
    ncols = len(headers)

    # 标题
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncols)
    ws.cell(row=1, column=1, value=title).font = s["title_font"]
    ws.cell(row=1, column=1).alignment = s["center"]
    ws.row_dimensions[1].height = 28

    # 表头
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=2, column=c, value=h)
        cell.font = s["header_font"]
        cell.fill = s["header_fill"]
        cell.alignment = s["center"]
        cell.border = s["border"]
    ws.row_dimensions[2].height = 22

    # 示例数据
    for r_idx, row in enumerate(sample_rows, start=3):
        for c_idx, val in enumerate(row, start=1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.font = s["normal_font"]
            cell.border = s["border"]
            cell.alignment = s["left_align"]
            if (r_idx - 3) % 2 == 1:
                cell.fill = s["alt_fill"]

    # 列宽
    for i, w in enumerate(col_widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    ws.freeze_panes = "A3"
    wb.save(str(path))


def _create_specific_plan_xlsx(path, client_name, audit_year):
    """具体审计计划汇总表"""
    headers = ["科目", "认定", "风险等级", "拟执行程序", "样本量", "执行人", "完成日期", "复核人", "备注"]
    sample = [
        ["货币资金", "存在/计价", "高", "余额调节表+双向核对+大额检查+截止测试", "100%", "", "", "", "对所有银行账户全面核查"],
        ["应收账款", "存在/计价", "高", "函证+账龄分析+坏账复核+截止测试", "MUS抽样", "", "", "", "重点关注长账龄"],
        ["存货", "存在/计价", "中", "盘点+跌价复核+截止测试", "重要项目+随机", "", "", "", ""],
        ["固定资产", "存在/计价", "中", "折旧重算+抽样盘点+减值测试", "MUS抽样", "", "", "", ""],
        ["收入", "完整性/截止", "高", "截止测试+本福特检验+分析性程序", "前后5天", "", "", "", "特别风险"],
        ["费用", "完整性/截止", "中", "截止测试+抽凭", "前后5天+25笔", "", "", "", ""],
        ["应付账款", "完整性", "中", "函证+账龄分析+截止测试", "MUS抽样", "", "", "", ""],
        ["在建工程", "存在/计价", "高", "明细检查+大额合同+资本化利息测试", "前10大", "", "", "", "如有"],
    ]
    widths = [14, 12, 10, 36, 14, 12, 12, 10, 22]
    _write_xlsx_template(
        path, "具体审计计划",
        f"{client_name} {audit_year}年度 — 具体审计计划汇总表",
        headers, sample, widths
    )


def _create_risk_assessment_xlsx(path, client_name, audit_year):
    """风险评估结果汇总表"""
    headers = ["风险编号", "风险描述", "涉及科目", "涉及认定", "是否特别风险", "拟应对措施", "执行结果"]
    sample = [
        ["R01", "收入舞弊（管理层凌驾于控制之上）", "营业收入", "完整性/发生", "是", "截止测试+分析性程序+本福特检验", ""],
        ["R02", "应收账款回收风险", "应收账款", "计价", "否", "账龄分析+函证+坏账复核", ""],
        ["R03", "[请填写]", "", "", "", "", ""],
    ]
    widths = [10, 36, 14, 18, 14, 36, 22]
    _write_xlsx_template(
        path, "风险评估",
        f"{client_name} {audit_year}年度 — 风险评估结果汇总表",
        headers, sample, widths
    )


def _create_group_materiality_xlsx(path, client_name, audit_year):
    """集团审计重要性分配表"""
    headers = ["组成部分名称", "重要性程度", "总资产占比", "营业收入占比", "组件重要性", "依据说明"]
    sample = [
        ["[母公司]", "重要", "100%", "100%", "[使用集团重要性]", "母公司"],
        ["[子公司A]", "重要", "30%", "25%", "≤0.8×集团重要性", "重要组成部分"],
        ["[子公司B]", "非重要", "5%", "3%", "无需单独重要性", "非重要组成部分"],
    ]
    widths = [22, 12, 14, 14, 18, 24]
    _write_xlsx_template(
        path, "重要性分配",
        f"{client_name} {audit_year}年度 — 集团审计重要性分配表",
        headers, sample, widths
    )


def _create_automation_log_xlsx(path, client_name, audit_year):
    """自动化工具使用记录（覆盖古立特全部工具，按使用频率排序）"""
    headers = ["序号", "工具名称", "版本", "用途", "数据来源", "数据验证步骤", "人工复核逻辑", "执行人", "执行日期"]
    sample = [
        [1, "古立特 MCP Server", f"v{_GV}", "审计计算工具集（35 个工具，含审计 13 + 数据 4 + 文档 6 + 财务分析 5 + 市场数据 4）", "", "", "AI 完成的匹配/计算结果由审计员复核", "", ""],
        # 底稿编制类（最常用）
        [2, "workpaper_init", f"v{_GV}", "底稿框架初始化", "客户名+审计年度", "目录结构生成校验", "模板内容人工填充", "", ""],
        [3, "materiality_calculator", f"v{_GV}", "重要性计算（CSA 1320）", "财务数据", "基准选择合理性", "百分比合理性人工评估", "", ""],
        [4, "analytical_procedures", f"v{_GV}", "分析性程序", "本期+上期对比表", "波动率计算正确性", "异常波动原因人工核实", "", ""],
        [5, "balance_sheet_process", f"v{_GV}", "科目余额表处理", "客户提供的TB", "检查重复值/格式校验", "末级科目正确性人工抽查", "", ""],
        [6, "bank_reconciliation", f"v{_GV}", "银行余额调节表", "对账单+日记账", "金额/日期对齐", "未达账项原因人工核实", "", ""],
        [7, "aging_analysis", f"v{_GV}", "账龄分析", "应收/应付明细", "对方名称去重", "长账龄项目逐笔核实", "", ""],
        [8, "depreciation_check", f"v{_GV}", "固定资产折旧重算", "固定资产明细表", "折旧方法识别正确性", "差异项原因人工核实", "", ""],
        [9, "audit_sampling", f"v{_GV}", "审计抽样", "总体数据", "总体完整性检查", "样本充分性人工评估", "", ""],
        [10, "cutoff_test", f"v{_GV}", "截止测试", "序时账", "日期格式校验", "可疑项凭证逐笔核实", "", ""],
        [11, "confirmation_letter", f"v{_GV}", "函证生成", "抽样样本", "对方名称/地址完整性", "函证内容人工最终确认", "", ""],
        [12, "audit_adjustment", f"v{_GV}", "三栏TB", "未审TB+调整分录", "借贷平衡校验", "调整分录依据人工核实", "", ""],
        [13, "benford_law_check", f"v{_GV}", "本福特定律检验", "财务数据序列", "样本量充分性", "异常分布原因人工分析", "", ""],
        [14, "address_split", f"v{_GV}", "地址拆分省市区", "客户/供应商地址", "拆分准确性", "异常地址人工核对", "", ""],
        # 文档处理类
        [15, "document_ocr", f"v{_GV}", "文档 OCR 识别", "PDF/图片", "MinerU 云端解析", "识别结果人工校对", "", ""],
        [16, "document_ocr_batch", f"v{_GV}", "批量 OCR", "多个 PDF/图片", "MinerU 云端解析", "识别结果人工校对", "", ""],
        [17, "pdf_merge", f"v{_GV}", "PDF 合并", "多个 PDF 文件", "页数完整性", "合并顺序人工确认", "", ""],
        [18, "pdf_split", f"v{_GV}", "PDF 拆分", "PDF 文件", "页码范围正确性", "拆分结果人工核对", "", ""],
        [19, "pdf_extract", f"v{_GV}", "PDF 提取页面", "PDF 文件", "页码正确性", "提取结果人工核对", "", ""],
        [20, "voucher_pdf_split", f"v{_GV}", "凭证 PDF 拆分", "凭证合订 PDF", "凭证号识别", "拆分结果人工核对", "", ""],
        # 数据分析类
        [21, "data_file_overview", f"v{_GV}", "数据概览", "Excel/CSV", "行列/类型/缺失值", "数据完整性人工判断", "", ""],
        [22, "data_group_summary", f"v{_GV}", "分组汇总", "Excel/CSV", "聚合函数正确性", "汇总结果合理性", "", ""],
        [23, "data_filter", f"v{_GV}", "数据筛选", "Excel/CSV", "筛选条件准确性", "筛选结果合理性", "", ""],
        [24, "data_pivot", f"v{_GV}", "透视表", "Excel/CSV", "维度选择合理性", "透视结果人工解读", "", ""],
        # 市场数据类
        [25, "stock_quote", f"v{_GV}", "A股实时行情", "akshare 接口", "数据时效性", "数据合理性人工判断", "", ""],
        [26, "stock_history", f"v{_GV}", "历史K线(A股/港股/美股)", "akshare 接口", "时间序列完整性", "异常波动人工分析", "", ""],
        [27, "stock_financial", f"v{_GV}", "A股财务数据", "新浪财经接口", "报告期完整性", "数据与公告核对", "", ""],
        [28, "report_download", f"v{_GV}", "上市公司公告下载", "巨潮资讯网", "文件完整性", "下载内容人工查验", "", ""],
        # 财务分析类
        [29, "reclassification", f"v{_GV}", "往来款重分类（CAS 30）", "往来明细表", "借贷方向校验", "调整分录人工核实", "", ""],
        [30, "cashflow_test", f"v{_GV}", "现金流量表测算（间接法）", "审定TB+利润表", "TB完整性检查", "测算结果人工核实", "", ""],
        [31, "cashflow_direct", f"v{_GV}", "现金流量表（直接法）", "序时账", "现金类科目筛选", "未归类项人工归类", "", ""],
        [32, "non_recurring_items", f"v{_GV}", "非经常性损益计算", "证监会项目清单", "项目分类正确性", "扣非净利润人工核对", "", ""],
        [33, "financial_ratios", f"v{_GV}", "财务比率分析", "资产负债表+利润表", "公式正确性", "比率合理性人工解读", "", ""],
        [34, "tax_adjustment", f"v{_GV}", "纳税调整明细", "科目余额表", "限额规则正确性", "需人工确认项逐项核实", "", ""],
        [35, "voucher_scan", f"v{_GV}", "凭证巡检（七大类风险）", "序时账", "可疑项识别规则", "可疑项逐笔核实", "", ""],
        [36, "fuzzy_match", f"v{_GV}", "名称模糊匹配", "两份名单", "匹配阈值合理性", "未匹配项人工审核", "", ""],
    ]
    widths = [6, 24, 10, 32, 18, 22, 26, 12, 12]
    _write_xlsx_template(
        path, "自动化工具记录",
        f"{client_name} {audit_year}年度 — 自动化工具使用记录（ISA 500 合规）",
        headers, sample, widths
    )


def _create_judgment_xlsx(path, client_name, audit_year):
    """重大职业判断记录单"""
    headers = ["判断事项", "判断内容", "判断依据", "支持证据", "判断结论", "判断人", "判断日期"]
    sample = [
        ["重要性水平的确定", "整体重要性 / 实际执行重要性 / 微小错报", "税前利润法 / 资产法 / 收入法的选择",
         "财务报表+经营稳定性分析", "[请填写]", "", ""],
        ["特别风险的识别", "如：收入确认、管理层凌驾", "风险评估程序结果", "询问记录+分析性程序结果",
         "[请填写]", "", ""],
        ["会计估计合理性", "如：坏账准备计提比例", "历史数据+行业对比", "账龄分析+减值测试结果",
         "[请填写]", "", ""],
        ["以前年度损益调整", "金额及性质", "管理层说明+原始凭证", "[支持文件]",
         "[请填写]", "", ""],
        ["持续经营能力评估", "12个月内是否存在重大不确定性", "现金流预测+经营计划+银行授信",
         "管理层书面说明", "[请填写]", "", ""],
    ]
    widths = [20, 26, 24, 22, 22, 12, 12]
    _write_xlsx_template(
        path, "职业判断",
        f"{client_name} {audit_year}年度 — 重大职业判断记录单",
        headers, sample, widths
    )


def _create_uncorrected_misstatement_xlsx(path, client_name, audit_year):
    """未更正错报汇总表"""
    headers = ["编号", "错报描述", "涉及科目", "影响金额（资产）", "影响金额（负债）", "影响金额（损益）", "管理层不调整理由", "审计师评价"]
    sample = [
        ["UM01", "[错报描述]", "[科目]", "", "", "", "[管理层理由]", "[影响评价]"],
    ]
    widths = [8, 32, 14, 16, 16, 16, 26, 22]
    _write_xlsx_template(
        path, "未更正错报",
        f"{client_name} {audit_year}年度 — 未更正错报汇总表",
        headers, sample, widths
    )


def _create_review_log_xlsx(path, client_name, audit_year):
    """复核记录表"""
    headers = ["复核级别", "复核人", "复核日期", "复核范围", "复核意见", "执行人回复", "回复日期", "状态"]
    sample = [
        ["项目组互审", "", "", "底稿完整性+程序充分性", "", "", "", "□待复核 □已通过"],
        ["项目经理复核", "", "", "全部底稿+审计结论", "", "", "", "□待复核 □已通过"],
        ["质量复核（如适用）", "", "", "重大判断+审计意见+披露", "", "", "", "□待复核 □已通过"],
    ]
    widths = [16, 12, 12, 24, 26, 22, 12, 16]
    _write_xlsx_template(
        path, "复核记录",
        f"{client_name} {audit_year}年度 — 项目组复核记录表",
        headers, sample, widths
    )


def _create_post_event_xlsx(path, client_name, audit_year):
    """期后事项复核表"""
    headers = ["事项类型", "事项描述", "发生日期", "金额影响", "是否调整事项", "处理方式", "支持证据", "复核人"]
    sample = [
        ["重大诉讼", "", "", "", "□调整 □披露 □无需处理", "", "", ""],
        ["资产减值", "", "", "", "", "", "", ""],
        ["重大销售退回", "", "", "", "", "", "", ""],
        ["管理层重大决议", "", "", "", "", "", "", ""],
    ]
    widths = [16, 32, 12, 16, 22, 16, 18, 12]
    _write_xlsx_template(
        path, "期后事项",
        f"{client_name} {audit_year}年度 — 期后事项复核表",
        headers, sample, widths
    )


def _create_workpaper_index(path, client_name, audit_year, folders, templates):
    """底稿索引（根目录）"""
    headers = ["索引号", "类别", "底稿名称", "存放位置", "状态", "执行人", "完成日期", "复核人"]

    sample = [
        ["A001", "计划与风险", "业务承接评价表", "1-计划与风险/业务承接评价表.docx", "□未开始 □进行中 □已完成", "", "", ""],
        ["A002", "计划与风险", "总体审计策略", "1-计划与风险/总体审计策略.docx", "", "", "", ""],
        ["A003", "计划与风险", "具体审计计划", "1-计划与风险/具体审计计划汇总表.xlsx", "", "", "", ""],
        ["A004", "计划与风险", "风险评估结果", "1-计划与风险/风险评估结果汇总表.xlsx", "", "", "", ""],
        ["B001", "实质性程序-货币资金", "银行存款余额调节表", "3-实质性程序/货币资金/", "", "", "", ""],
        ["B002", "实质性程序-货币资金", "双向核对结果", "3-实质性程序/货币资金/", "", "", "", ""],
        ["B003", "实质性程序-货币资金", "截止测试", "3-实质性程序/货币资金/", "", "", "", ""],
        ["B011", "实质性程序-应收", "账龄分析", "3-实质性程序/往来款项/", "", "", "", ""],
        ["B012", "实质性程序-应收", "函证清单", "3-实质性程序/函证/", "", "", "", ""],
        ["B021", "实质性程序-固定资产", "折旧重算", "3-实质性程序/[扩展]/", "", "", "", ""],
        ["B031", "实质性程序-收入", "截止测试", "3-实质性程序/损益类/", "", "", "", ""],
        ["B032", "实质性程序-收入", "本福特检验", "3-实质性程序/损益类/", "", "", "", ""],
        ["C001", "准则合规", "自动化工具使用记录", "4-准则合规/自动化工具使用记录.xlsx", "", "", "", ""],
        ["D001", "职业判断", "重大职业判断记录单", "5-职业判断/重大职业判断记录单.xlsx", "", "", "", ""],
        ["E001", "完成阶段", "审计调整分录汇总", "6-完成阶段/", "", "", "", ""],
        ["E002", "完成阶段", "未更正错报汇总", "6-完成阶段/未更正错报汇总表.xlsx", "", "", "", ""],
        ["E003", "完成阶段", "审计总结", "6-完成阶段/审计总结.docx", "", "", "", ""],
        ["E004", "完成阶段", "复核记录", "6-完成阶段/复核记录表.xlsx", "", "", "", ""],
        ["E005", "完成阶段", "期后事项复核", "6-完成阶段/期后事项复核表.xlsx", "", "", "", ""],
        ["E006", "完成阶段", "审计报告草稿", "6-完成阶段/审计报告（草稿）.docx", "", "", "", ""],
    ]
    widths = [8, 22, 22, 36, 24, 12, 12, 12]
    _write_xlsx_template(
        path, "底稿索引",
        f"{client_name} {audit_year}年度审计底稿 — 索引（古立特生成）",
        headers, sample, widths
    )


def _create_readme(path, client_name, audit_year):
    """根目录 README"""
    content = f"""═══════════════════════════════════════════════════════════
  {client_name} {audit_year} 年度审计底稿
  由古立特 Gridman v{_GV} 自动生成
═══════════════════════════════════════════════════════════

【目录结构】

  1-计划与风险/      审计计划、风险评估
  2-控制测试/        控制测试底稿（仅当依赖控制时）
  3-实质性程序/      各科目实质性程序底稿
  4-准则合规/        自动化工具使用记录（ISA 500）
  5-职业判断/        重大职业判断记录
  6-完成阶段/        汇总、总结、报告草稿
  7-备查类/          支持性文件、合同、章程
  底稿索引.xlsx      所有底稿索引

【配色说明】

  古立特生成的所有 Excel 模板使用统一灰紫配色：
    深灰紫 #2D2B55  — 表头背景
    暗灰紫 #3D3A6B  — 合计行
    浅灰紫 #EEEAF5  — 小计行
    极浅灰紫 #F7F5FA — 隔行变色
    边框 #D4D0E0
    字体 微软雅黑

  ⚠ 配色可改：
  事务所如有内部配色规范，可手动调整。古立特只是提供
  默认配色，不强制统一。

【封面/抬头/装订说明】

  古立特只生成"核心内容"。事务所内部的封面、抬头格式、
  装订规范、签字栏样式等，由用户根据所内规定自行加上。

【模板使用顺序】

  Step 1: 业务承接评价表 → 决定是否承接
  Step 2: 用 materiality_calculator 计算重要性
  Step 3: 总体审计策略 → 填入重要性、范围、时间安排
  Step 4: 风险评估结果汇总 → 识别风险并设计应对
  Step 5: 具体审计计划 → 各科目程序设计
  Step 6: 各科目实质性程序（用古立特工具跑）
  Step 7: 自动化工具使用记录（每用一次工具就记录）
  Step 8: 重大职业判断记录（关键判断必填）
  Step 9: 调整分录汇总 + 三栏TB（用 audit_adjustment）
  Step 10: 未更正错报、复核记录、期后事项
  Step 11: 审计总结 + 审计报告草稿

【AI 与人工分工】

  ✓ AI 能做：
    - 数据计算（调节、账龄、折旧、抽样、截止）
    - 模板填充
    - 异常项标记
    - 报表勾稽校验

  ✗ AI 不能做（必须人工）：
    - 实地盘点
    - 询问管理层
    - 检查实物凭证原件
    - 评价管理层诚信
    - 形成最终审计意见
    - 签字确认

═══════════════════════════════════════════════════════════
  生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
═══════════════════════════════════════════════════════════
"""
    path.write_text(content, encoding="utf-8")
